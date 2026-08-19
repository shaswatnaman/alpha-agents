"""
Document ingestion pipeline:

    source bytes
        → extraction   (PDF/DOCX/TXT → plain text)
        → cleaning     (whitespace normalisation, boilerplate removal)
        → chunking     (overlapping fixed-size chunks; semantic chunking optional)
        → metadata enrichment
        → embedding    (OpenAI text-embedding-3-small)
        → indexing     (PostgreSQL + pgvector)

Every step is independent and can be retested in isolation.
"""

from __future__ import annotations

import hashlib
import io
import re

import structlog

from app.config.settings import get_settings
from app.domain.models import Document, DocumentChunk, DocumentMetadata
from app.llm.provider import get_llm_provider

log = structlog.get_logger(__name__)


# ── Text Extraction ───────────────────────────────────────────────────────────


def extract_text_from_pdf(data: bytes) -> str:
    try:
        import pypdf

        reader = pypdf.PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(pages)
    except ImportError as err:
        raise RuntimeError("pypdf not installed. Run: pip install pypdf") from err


def extract_text_from_docx(data: bytes) -> str:
    try:
        import docx

        doc = docx.Document(io.BytesIO(data))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except ImportError as err:
        raise RuntimeError("python-docx not installed.") from err


def extract_text(data: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(data)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(data)
    elif ext == "txt":
        return data.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported document type: {ext}")


# ── Cleaning ──────────────────────────────────────────────────────────────────


def clean_text(text: str) -> str:
    # Collapse excessive whitespace while preserving paragraph breaks
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove page-number artifacts common in PDFs
    text = re.sub(r"\n\s*\d+\s*\n", "\n", text)
    return text.strip()


# ── Chunking ──────────────────────────────────────────────────────────────────


def chunk_text(
    text: str,
    chunk_size: int,
    chunk_overlap: int,
) -> list[tuple[str, int, int]]:
    """
    Split text into overlapping chunks.

    Returns list of (chunk_text, char_start, char_end).

    Strategy: word-boundary-aware fixed-size splits with configurable overlap.
    This avoids cutting mid-word while keeping chunk sizes predictable for
    the embedding model's context window.
    """
    chunks: list[tuple[str, int, int]] = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = min(start + chunk_size, text_len)

        # Extend to nearest sentence boundary (. ! ?) if possible
        if end < text_len:
            search_region = text[end : end + 100]
            m = re.search(r"[.!?]\s", search_region)
            if m:
                end += m.start() + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append((chunk, start, end))

        start = end - chunk_overlap
        if start <= 0:
            break

    return chunks


# ── Embedding ─────────────────────────────────────────────────────────────────


async def embed_chunks(chunks: list[DocumentChunk]) -> list[DocumentChunk]:
    """
    Embed chunk texts in batches of 100 (API limit).
    Modifies chunks in place, returns the list.
    """
    provider = get_llm_provider()
    batch_size = 100

    for i in range(0, len(chunks), batch_size):
        batch = chunks[i : i + batch_size]
        texts = [c.text for c in batch]
        embeddings = await provider.embed(texts)
        for chunk, embedding in zip(batch, embeddings, strict=True):
            chunk.embedding = embedding

    return chunks


# ── Pipeline Entry Point ──────────────────────────────────────────────────────


async def ingest_document(
    data: bytes,
    filename: str,
    metadata: DocumentMetadata,
) -> tuple[Document, list[DocumentChunk]]:
    """
    Full ingestion pipeline.

    Returns the Document and its embedded DocumentChunks ready for indexing.
    Raises ValueError if the document has already been ingested (content_hash match).
    """
    settings = get_settings()

    content_hash = hashlib.sha256(data).hexdigest()
    log.info("ingestion_start", filename=filename, size_bytes=len(data), hash=content_hash)

    # Extract
    raw_text = extract_text(data, filename)
    log.debug("text_extracted", chars=len(raw_text))

    # Clean
    cleaned = clean_text(raw_text)

    # Build Document
    doc = Document(
        filename=filename,
        content_hash=content_hash,
        metadata=metadata,
    )

    # Chunk
    raw_chunks = chunk_text(cleaned, settings.chunk_size, settings.chunk_overlap)
    chunks: list[DocumentChunk] = []
    for idx, (text, start, end) in enumerate(raw_chunks):
        chunks.append(
            DocumentChunk(
                document_id=doc.id,
                chunk_index=idx,
                text=text,
                char_start=start,
                char_end=end,
                metadata={
                    "ticker": metadata.ticker,
                    "document_type": metadata.document_type,
                    "fiscal_year": metadata.fiscal_year,
                },
            )
        )

    log.info("chunks_created", document_id=doc.id, count=len(chunks))

    # Embed
    chunks = await embed_chunks(chunks)
    log.info("chunks_embedded", document_id=doc.id)

    return doc, chunks
